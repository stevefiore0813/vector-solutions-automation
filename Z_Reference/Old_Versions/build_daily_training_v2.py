import argparse, datetime as dt, hashlib, json, logging, os, random, re
from pathlib import Path
import yaml

try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

FM_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)
TITLE_RE = re.compile(r"^\s*(module|training|lesson)\s*:\s*(.+)$", re.IGNORECASE)

SITE_TRAINING_TYPES = {
    # Keep this list synced to the site’s labels exactly as shown in the checkbox group
    "Building Construction",
    "Fire Behavior",
    "Fire Detection, Alarm Systems, Suppression Systems",
    "Fire Extinguishers",
    "Fire Hose Evolutions",
    "Fire Streams and Nozzles",
    "Forcible Entry",
    "Ground Ladders",
    "Salvage and Overhaul",
    "Search and Rescue",
    "Ventilation",
    "VES - Ventilation, Enter, Search",
    "Water Supply (ex. hydrant operations, tender operations, dry hydrants)",
    "Firefighting Tactics and Strategies",
    "Extrication"
}

def load_cfg(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def init_logging(log_dir: Path):
    log_dir.mkdir(parents=True, exist_ok=True)
    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    logfile = log_dir / f"build_daily_training_v2_{stamp}.log"
    logging.basicConfig(level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=[logging.FileHandler(logfile, encoding="utf-8"), logging.StreamHandler()])
    return logfile

def date_seed(fixed=None):
    if fixed is not None: return int(fixed)
    today = dt.date.today().isoformat()
    return int(hashlib.sha256(today.encode()).hexdigest(), 16) % (2**31 - 1)

def parse_front_matter(text: str) -> tuple[dict, str]:
    m = FM_RE.match(text)
    if not m:
        return {}, text
    meta = yaml.safe_load(m.group(1)) or {}
    rest = text[m.end():]
    return meta, rest

def split_md_modules(text: str):
    blocks = []
    cur_title, cur_body = None, []
    for line in text.splitlines():
        if line.lstrip().startswith("#"):
            heading = re.sub(r"^#+\s*", "", line).strip()
            mt = TITLE_RE.match(heading)
            if mt:
                if cur_title:
                    blocks.append((cur_title.strip(), "\n".join(cur_body).strip()))
                cur_title, cur_body = mt.group(2), []
                continue
        if cur_title:
            cur_body.append(line)
    if cur_title:
        blocks.append((cur_title.strip(), "\n".join(cur_body).strip()))
    return blocks

def parse_md(path: Path):
    text = path.read_text(encoding="utf-8")
    meta, rest = parse_front_matter(text)
    modules = []
    for title, body in split_md_modules(rest if rest.strip() else text):
        modules.append({"title": title, "meta": meta.copy(), "body": body.strip()})
    logging.info(f"MD {path.name}: {len(modules)} modules")
    return modules

def parse_docx(path: Path):
    if not HAVE_DOCX:
        logging.warning("python-docx missing; skipping %s", path.name)
        return []
    doc = Document(str(path))
    raw = "\n".join(p.text for p in doc.paragraphs)
    meta, _ = parse_front_matter(raw)
    # build module list scanning headings
    modules, cur_title, cur = [], None, []
    def is_heading(p):
        try: return p.style and 'heading' in p.style.name.lower()
        except: return False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: 
            if cur_title: cur.append("")
            continue
        if is_heading(p):
            mt = TITLE_RE.match(t)
            if mt:
                if cur_title:
                    modules.append({"title": cur_title.strip(), "meta": meta.copy(), "body": "\n".join(cur).strip()})
                cur_title, cur = mt.group(2), []
                continue
        if cur_title:
            cur.append(t)
    if cur_title:
        modules.append({"title": cur_title.strip(), "meta": meta.copy(), "body": "\n".join(cur).strip()})
    logging.info(f"DOCX {path.name}: {len(modules)} modules")
    return modules

def discover(input_dir: Path):
    allm = []
    for p in sorted(input_dir.glob("**/*")):
        if p.is_dir(): continue
        if p.suffix.lower() == ".md": allm += parse_md(p)
        elif p.suffix.lower() == ".docx": allm += parse_docx(p)
    return allm

def normalize_types(vals):
    picked = set()
    for v in vals or []:
        v = str(v).strip()
        # exact match preferred
        if v in SITE_TRAINING_TYPES:
            picked.add(v); continue
        # loose contains match fallback
        for label in SITE_TRAINING_TYPES:
            if v.lower() in label.lower() or label.lower() in v.lower():
                picked.add(label)
    return sorted(picked)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="05_Dev_Env/Dependencies/config.yaml")
    ap.add_argument("--date", default=None)
    ap.add_argument("--location", default=None)
    ap.add_argument("--duration", type=float, default=None)
    ap.add_argument("--instructor", default=None)
    args = ap.parse_args()

    cfg = load_cfg(Path(args.config))
    input_dir = Path(cfg["paths"]["input_dir"])
    out_dir = Path(cfg["paths"]["output_dir"])
    log_dir = Path(cfg["paths"]["log_dir"])
    logfile = init_logging(log_dir)

    date_str = args.date or dt.date.today().isoformat()
    rng = random.Random(date_seed(cfg.get("fixed_seed")))
    modules = discover(input_dir)
    k = int(cfg.get("modules_per_day", 3))
    if not modules:
        logging.error("No modules found.")
        return 2
    chosen = modules if k >= len(modules) else rng.sample(modules, k)

    # Compose description by concatenating bodies with headings
    desc_lines = [f"Daily Training - {date_str}", ""]
    titles = []
    collected_types = set()
    duration_tot = 0.0
    instr = None
    for m in chosen:
        titles.append(m["title"])
        desc_lines.append(f"## {m['title']}")
        if m["body"]:
            desc_lines.append("")
            desc_lines.append(m["body"])
            desc_lines.append("")
        meta = m.get("meta") or {}
        collected_types |= set(normalize_types(meta.get("training_types")))
        if meta.get("duration_hours"):
            try: duration_tot += float(meta["duration_hours"])
            except: pass
        if meta.get("instructor"): instr = meta["instructor"]

    if args.duration is not None:
        duration_tot = float(args.duration)
    elif duration_tot == 0:
        duration_tot = float(cfg.get("default_duration_hours", 1))

    payload = {
        "location": args.location or cfg.get("default_location", ""),
        "training_types": sorted(collected_types) or cfg.get("default_training_types", []),
        "description": "\n".join(desc_lines).strip() + "\n",
        "duration_hours": duration_tot,
        "date_complete": date_str,
        "instructor": args.instructor or instr or cfg.get("default_instructor", ""),
        "modules_included": titles
    }

    out_dir.mkdir(parents=True, exist_ok=True)
    md_name = cfg.get("output_name_pattern", "daily_training_{date}.md").replace("{date}", date_str)
    json_name = md_name.replace(".md", ".json")
    (out_dir / md_name).write_text(payload["description"], encoding="utf-8")
    (out_dir / json_name).write_text(json.dumps(payload, indent=2), encoding="utf-8")
    logging.info("Wrote: %s and %s", md_name, json_name)
    print(str(out_dir / json_name))
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
